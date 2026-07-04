import re
from difflib import SequenceMatcher

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from docx_utils import _iter_block_items
from schemas import InstructionItem


def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip().lstrip("@").strip())


def iter_all_paragraphs(doc: Document):
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            yield block
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def find_marker_spans(paragraphs: list[Paragraph]) -> list[dict]:
    spans: list[dict] = []
    index = 0
    while index < len(paragraphs):
        text = paragraphs[index].text
        if "@" not in text:
            index += 1
            continue

        start = index
        combined = text
        while "@@" not in combined and index + 1 < len(paragraphs):
            index += 1
            combined += "\n" + paragraphs[index].text

        if "@@" not in combined:
            index += 1
            continue

        match = re.search(r"@(.*?)(@@)", combined, re.DOTALL)
        if not match:
            index += 1
            continue

        inner = match.group(1).strip()
        before = combined[: match.start()]
        after = combined[match.end() :]
        spans.append(
            {
                "span_index": len(spans),
                "start": start,
                "end": index,
                "inner": inner,
                "before": before,
                "after": after,
                "inline": bool(before.strip() or after.strip()),
            }
        )
        index += 1
    return spans


def match_tag(inner_text: str, instructions: list[InstructionItem]) -> InstructionItem | None:
    norm_inner = _normalize(inner_text)
    best: InstructionItem | None = None
    best_score = 0.0

    for tag in instructions:
        norm_instruction = _normalize(tag.instruction)
        score = SequenceMatcher(None, norm_inner, norm_instruction).ratio()
        if norm_inner in norm_instruction or norm_instruction in norm_inner:
            score = max(score, 0.9)
        if score > best_score:
            best_score = score
            best = tag

    return best if best_score >= 0.5 else None


def apply_replacements_to_docx(
    template_path: str,
    output_path: str,
    spans: list[dict],
    replacements: dict[int, str],
) -> str:
    doc = Document(template_path)
    paragraphs = list(iter_all_paragraphs(doc))

    for span in spans:
        replacement = replacements.get(span["span_index"])
        if replacement is None:
            continue

        new_text = f"{span['before']}{replacement}{span['after']}"
        set_paragraph_text(paragraphs[span["start"]], new_text)
        for idx in range(span["start"] + 1, span["end"] + 1):
            set_paragraph_text(paragraphs[idx], "")

    doc.save(output_path)
    return output_path
