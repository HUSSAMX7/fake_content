from __future__ import annotations

import tempfile
import unittest
import zipfile
from pathlib import Path

from agents.chart_agent import prepare_charts_for_replacements
from agents.integrator import _enforce_chart_blocks
from docx import Document
from docx_skill_integrator import apply_replacements_to_docx, prepare_template
from schemas import ContentBlock, SpanWritingPlan


class ChartAgentTests(unittest.TestCase):
    def test_span_writing_plan_accepts_chart_mode(self) -> None:
        plan = SpanWritingPlan(mode="chart", items=[])
        self.assertEqual(plan.mode, "chart")

    def test_enforce_chart_keeps_intro_and_chart(self) -> None:
        blocks = [
            ContentBlock(type="paragraph", text="يوضح الشكل التالي توزيع الكميات:"),
            ContentBlock(
                type="chart",
                text="شكل (1): توزيع الكميات",
                chart_kind="bar",
                chart_title="توزيع الكميات",
                chart_labels=["أ", "ب", "ج"],
                chart_values=[10.0, 20.0, 30.0],
            ),
            ContentBlock(type="heading", text="تفاصيل إضافية"),
            ContentBlock(type="bullet_item", text="لا يجب أن يبقى"),
        ]
        enforced = _enforce_chart_blocks(blocks)
        self.assertEqual(len(enforced), 2)
        self.assertEqual(enforced[0].type, "paragraph")
        self.assertEqual(enforced[1].type, "chart")

    def test_prepare_keeps_valid_chart_block(self) -> None:
        replacements = {
            0: [
                ContentBlock(
                    type="chart",
                    text="شكل (1): الكميات",
                    chart_kind="bar",
                    chart_title="الكميات",
                    chart_labels=["مرحلة أ", "مرحلة ب"],
                    chart_values=[12.0, 8.0],
                )
            ]
        }
        updated = prepare_charts_for_replacements(replacements)
        self.assertEqual(updated[0][0].type, "chart")
        self.assertEqual(updated[0][0].chart_kind, "bar")

    def test_soft_fail_missing_values(self) -> None:
        replacements = {
            2: [
                ContentBlock(
                    type="chart",
                    text="تعليق احتياطي",
                    chart_kind="bar",
                    chart_labels=["أ"],
                    chart_values=None,
                )
            ]
        }
        updated = prepare_charts_for_replacements(replacements)
        self.assertEqual(updated[2][0].type, "paragraph")
        self.assertEqual(updated[2][0].text, "تعليق احتياطي")

    def test_native_bar_chart_embedded_in_docx(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            output = root / "out.docx"
            doc = Document()
            doc.add_paragraph(
                "@أدرج سمارت شارت أعمدة للكميات من المصادر فقط@@"
            )
            doc.save(template)

            temp_dir, unpacked_dir, spans = prepare_template(str(template))
            try:
                self.assertEqual(len(spans), 1)
                apply_replacements_to_docx(
                    str(template),
                    str(output),
                    spans,
                    {
                        spans[0]["span_index"]: [
                            ContentBlock(
                                type="chart",
                                text="شكل (1): الكميات",
                                chart_kind="bar",
                                chart_title="الكميات",
                                chart_labels=["أ", "ب", "ج"],
                                chart_values=[10.0, 20.0, 5.0],
                            )
                        ]
                    },
                    unpacked_dir=unpacked_dir,
                )
            finally:
                temp_dir.cleanup()

            with zipfile.ZipFile(output) as archive:
                names = archive.namelist()
                chart_parts = [name for name in names if name.startswith("word/charts/chart")]
                embedding_parts = [
                    name for name in names if name.startswith("word/embeddings/") and name.endswith(".xlsx")
                ]
                self.assertTrue(chart_parts, names)
                self.assertTrue(embedding_parts, names)
                document_xml = archive.read("word/document.xml").decode("utf-8")
                self.assertIn("drawingml/2006/chart", document_xml)
                self.assertIn("شكل (1): الكميات", document_xml)

    def test_flow_chart_becomes_editable_table(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            output = root / "out.docx"
            doc = Document()
            doc.add_paragraph("@أدرج سمارت شارت flow لمراحل المشروع@@")
            doc.save(template)

            temp_dir, unpacked_dir, spans = prepare_template(str(template))
            try:
                apply_replacements_to_docx(
                    str(template),
                    str(output),
                    spans,
                    {
                        spans[0]["span_index"]: [
                            ContentBlock(
                                type="chart",
                                text="شكل (1): المراحل",
                                chart_kind="flow",
                                chart_labels=["التجهيز", "التصميم", "التنفيذ"],
                            )
                        ]
                    },
                    unpacked_dir=unpacked_dir,
                )
            finally:
                temp_dir.cleanup()

            with zipfile.ZipFile(output) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
                self.assertIn("<w:tbl", document_xml)
                self.assertIn("التجهيز", document_xml)
                self.assertIn("التنفيذ", document_xml)


if __name__ == "__main__":
    unittest.main()
